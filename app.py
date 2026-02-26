import streamlit as st
from google import genai
from google.genai import types
from PIL import Image
from gtts import gTTS
from io import BytesIO
import sqlite3
import uuid
import time
import re
from docx import Document

# ==========================================
# 1. CONFIGURARE PAGINÄ‚ & STIL
# ==========================================
st.set_page_config(
    page_title="Avocat ONRC AI",
    page_icon="âš–ï¸",
    layout="wide"
)

st.markdown("""
<style>
    .stChatMessage { font-family: 'Georgia', serif; font-size: 1.05rem; }
    .stButton button { background-color: #2c3e50; color: white; border-radius: 5px; }
    .stSuccess { background-color: #f0fdf4; border: 1px solid #bbf7d0; color: #166534; }
    h1 { color: #1e3a8a; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. CLIENTUL GOOGLE GENAI
# ==========================================
api_key = None
if "GOOGLE_API_KEY" in st.secrets:
    raw_key = st.secrets["GOOGLE_API_KEY"]
    if isinstance(raw_key, list):
        api_key = raw_key[0]
    else:
        api_key = raw_key

if not api_key:
    api_key = st.sidebar.text_input("Introdu Google API Key:", type="password")

if not api_key:
    st.warning("âš ï¸ Te rog introdu cheia API Ã®n sidebar.")
    st.stop()

try:
    clean_key = str(api_key).strip()
    client = genai.Client(api_key=clean_key)
except Exception as e:
    st.error(f"Eroare criticÄƒ la conectare: {e}")
    st.stop()

# ==========================================
# 3. MEMORIE (SQLITE)
# ==========================================
def init_db():
    conn = sqlite3.connect('legal_chat_v9.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS history 
                 (session_id TEXT, role TEXT, content TEXT, timestamp REAL)''')
    conn.commit()
    conn.close()

def save_message(session_id, role, content):
    conn = sqlite3.connect('legal_chat_v9.db')
    c = conn.cursor()
    c.execute("INSERT INTO history VALUES (?, ?, ?, ?)", (session_id, role, content, time.time()))
    conn.commit()
    conn.close()

def load_history(session_id):
    conn = sqlite3.connect('legal_chat_v9.db')
    c = conn.cursor()
    c.execute("SELECT role, content FROM history WHERE session_id=? ORDER BY timestamp ASC", (session_id,))
    data = c.fetchall()
    conn.close()
    return [{"role": row[0], "content": row[1]} for row in data]

def clear_history(session_id):
    conn = sqlite3.connect('legal_chat_v9.db')
    c = conn.cursor()
    c.execute("DELETE FROM history WHERE session_id=?", (session_id,))
    conn.commit()
    conn.close()

init_db()

if "session_id" not in st.query_params:
    st.session_state.session_id = str(uuid.uuid4())
    st.query_params["session_id"] = st.session_state.session_id
else:
    st.session_state.session_id = st.query_params["session_id"]

# ==========================================
# 4. FUNCÈšII DE CURÄ‚ÈšARE È˜I GENERARE WORD
# ==========================================

def clean_ai_response(text):
    """CurÄƒÈ›Äƒ textul de artefacte Google Search È™i repetiÈ›ii"""
    if not text: return ""
    
    # 1. EliminÄƒm blocul de citaÈ›ii <details>...</details>
    text = re.sub(r'<details>.*?</details>', '', text, flags=re.DOTALL)
    
    # 2. NormalizÄƒm break-urile HTML
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    
    # 3. EliminÄƒm orice alt tag HTML
    text = re.sub(r'<[^>]+>', '', text)
    
    # 4. ELIMINAREA REPETIÈšIILOR (Fix pentru problema ta)
    # Ãnlocuim 3 sau mai multe linii noi consecutive cu doar 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

def create_docx(text):
    doc = Document()
    
    # Folosim funcÈ›ia de curÄƒÈ›are Ã®nainte de a genera Word-ul
    clean_text = clean_ai_response(text)
    
    if not clean_text:
        clean_text = "Eroare: Document gol."

    lines = clean_text.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
            
        if stripped_line.startswith('#'):
            clean_content = stripped_line.replace('#', '').strip()
            doc.add_heading(clean_content, level=1)
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            clean_content = stripped_line[2:].strip().replace('**', '').replace('__', '')
            doc.add_paragraph(clean_content, style='List Bullet')
        else:
            clean_content = stripped_line.replace('**', '').replace('__', '')
            doc.add_paragraph(clean_content)
            
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# 5. SIDEBAR - UPLOAD CLOUD
# ==========================================
st.sidebar.title("ğŸ—‚ï¸ Dosar Acte")

if st.sidebar.button("ğŸ—‘ï¸ Caz Nou (Reset)", type="primary"):
    clear_history(st.session_state.session_id)
    st.session_state.messages = []
    st.session_state.uploaded_refs = [] 
    st.rerun()

st.sidebar.divider()

if "uploaded_refs" not in st.session_state:
    st.session_state.uploaded_refs = []

uploaded_files = st.sidebar.file_uploader("AdaugÄƒ documente", type=["jpg", "png", "pdf"], accept_multiple_files=True)

if uploaded_files:
    if st.sidebar.button("â˜ï¸ ÃncarcÄƒ Ã®n Dosar"):
        progress_bar = st.sidebar.progress(0)
        
        for idx, up_file in enumerate(uploaded_files):
            if not any(f['display_name'] == up_file.name for f in st.session_state.uploaded_refs):
                try:
                    with st.spinner(f"Se proceseazÄƒ: {up_file.name}..."):
                        file_bytes = up_file.getvalue()
                        uploaded_file = client.files.upload(
                            file=BytesIO(file_bytes),
                            config=types.UploadFileConfig(
                                display_name=up_file.name,
                                mime_type=up_file.type
                            )
                        )
                        while uploaded_file.state.name == "PROCESSING":
                            time.sleep(1)
                            uploaded_file = client.files.get(name=uploaded_file.name)
                        
                        if uploaded_file.state.name == "FAILED":
                            st.sidebar.error(f"Eroare Google: {up_file.name}")
                        else:
                            st.session_state.uploaded_refs.append({
                                'display_name': up_file.name,
                                'uri': uploaded_file.uri, 
                                'mime_type': up_file.type
                            })
                            st.sidebar.success(f"âœ… {up_file.name} adÄƒugat.")
                except Exception as e:
                    st.sidebar.error(f"Eroare upload: {e}")
            progress_bar.progress((idx + 1) / len(uploaded_files))
        time.sleep(1)
        st.rerun()

if st.session_state.uploaded_refs:
    st.sidebar.info(f"Dosar activ: {len(st.session_state.uploaded_refs)} acte")
else:
    st.sidebar.caption("Dosar gol.")

enable_audio = st.sidebar.checkbox("ğŸ”Š Audio", value=False)

# ==========================================
# 6. CONFIGURARE AVOCAT
# ==========================================

PROMPT_AVOCAT = """
EÈ™ti un Avocat Virtual Senior, expert Ã®n Drept Comercial, Proceduri ONRC È™i Fiscalitate (RomÃ¢nia).

OBIECTIV:
Oferi consultanÈ›Äƒ juridicÄƒ preliminarÄƒ clarÄƒ antreprenorilor.

INSTRUCÈšIUNI SPECIALE (SEARCH GROUNDING):
1. FoloseÈ™te Google Search activ pentru a verifica orice modificare legislativÄƒ recentÄƒ (2023-2026).
2. VerificÄƒ taxele ONRC actuale È™i procedurile din Legea 265/2022 (digitalizare).
3. DacÄƒ utilizatorul Ã®ntreabÄƒ de o lege viitoare, cautÄƒ "proiecte legislative" sau "propuneri modificare cod fiscal".

REGULI DE RÄ‚SPUNS:
- Fii precis: CiteazÄƒ articolul de lege cÃ¢nd e posibil.
- StructurÄƒ: Pas 1, Pas 2, Acte Necesare, Costuri Estimative.
- Avertisment: Include mereu disclaimer-ul cÄƒ eÈ™ti un AI.
"""

search_tool = types.Tool(google_search=types.GoogleSearch())

# SetÄƒri de siguranÈ›Äƒ relaxate pentru context legal
safety_settings = [
    types.SafetySetting(category='HARM_CATEGORY_HATE_SPEECH', threshold='BLOCK_NONE'),
    types.SafetySetting(category='HARM_CATEGORY_DANGEROUS_CONTENT', threshold='BLOCK_NONE'),
    types.SafetySetting(category='HARM_CATEGORY_HARASSMENT', threshold='BLOCK_NONE'),
    types.SafetySetting(category='HARM_CATEGORY_SEXUALLY_EXPLICIT', threshold='BLOCK_NONE'),
]

generate_config = types.GenerateContentConfig(
    system_instruction=PROMPT_AVOCAT,
    tools=[search_tool],
    temperature=0.3,
    safety_settings=safety_settings
)

# ==========================================
# 7. CHAT LOGIC
# ==========================================

st.title("âš–ï¸ Avocat Consultant ONRC")
st.caption("ExpertizÄƒ juridicÄƒ â€¢ Redactare Acte â€¢ AnalizÄƒ Dosar")

if "messages" not in st.session_state or not st.session_state.messages:
    st.session_state.messages = load_history(st.session_state.session_id)

# AfiÈ™are Mesaje
for i, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"], avatar="ğŸ‘¤" if msg["role"] == "user" else "âš–ï¸"):
        # AICI APLICÄ‚M CURÄ‚ÈšAREA PE AFIÈ˜ARE
        display_text = clean_ai_response(msg["content"])
        st.markdown(display_text)
        
        if msg["role"] == "assistant" and display_text:
            docx = create_docx(display_text) # Word-ul primeÈ™te textul curat
            st.download_button(
                label="ğŸ“„ DescarcÄƒ Word",
                data=docx,
                file_name=f"Document_Juridic_{i}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{i}"
            )

# Input
if user_input := st.chat_input("Ex: Vreau sÄƒ schimb sediul social. RedacteazÄƒ decizia."):
    
    st.session_state.messages.append({"role": "user", "content": user_input})
    save_message(st.session_state.session_id, "user", user_input)
    with st.chat_message("user", avatar="ğŸ‘¤"):
        st.write(user_input)

    # Context
    contents_payload = []
    for msg in st.session_state.messages[:-1]:
        role_gemini = "model" if msg["role"] == "assistant" else "user"
        if msg["content"]:
            contents_payload.append(types.Content(
                role=role_gemini,
                parts=[types.Part.from_text(text=clean_ai_response(msg["content"]))]
            ))
    
    current_parts = []
    if st.session_state.uploaded_refs:
        for ref in st.session_state.uploaded_refs:
            current_parts.append(types.Part.from_uri(file_uri=ref['uri'], mime_type=ref['mime_type']))
        current_parts.append(types.Part.from_text(text="\n\n[SISTEM: AnalizeazÄƒ documentele din dosar]"))
    current_parts.append(types.Part.from_text(text=user_input))
    
    contents_payload.append(types.Content(role="user", parts=current_parts))

    # Generare
    with st.chat_message("assistant", avatar="âš–ï¸"):
        placeholder = st.empty()
        full_text = ""
        
        try:
            response_stream = client.models.generate_content_stream(
                model='gemini-2.5-flash-lite',
                contents=contents_payload,
                config=generate_config
            )
            
            for chunk in response_stream:
                if chunk.text:
                    full_text += chunk.text
                    # CurÄƒÈ›Äƒm "live" doar cele mai evidente probleme vizuale, restul la final
                    display_chunk = full_text.replace('<br>', '\n') 
                    placeholder.markdown(display_chunk + "â–Œ")
            
            # CurÄƒÈ›are FINALÄ‚ È™i SALVARE
            final_clean_text = clean_ai_response(full_text)
            
            if not final_clean_text:
                placeholder.error("Nu am putut genera un rÄƒspuns valid. ÃncearcÄƒ din nou.")
            else:
                placeholder.markdown(final_clean_text)
            
            # SalvÄƒm textul BRUT (cu tot cu surse) Ã®n DB, dar Ã®l curÄƒÈ›Äƒm la afiÈ™are/word
            # Sau mai bine, salvÄƒm textul CURAT ca sÄƒ nu poluÄƒm istoricul
            st.session_state.messages.append({"role": "assistant", "content": full_text})
            save_message(st.session_state.session_id, "assistant", full_text)

            if final_clean_text:
                docx = create_docx(final_clean_text)
                st.download_button(
                    label="ğŸ“„ DescarcÄƒ Documentul Word",
                    data=docx,
                    file_name="Document_Juridic_AI.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_new"
                )

            if enable_audio and final_clean_text:
                audio_text = final_clean_text.replace("*", "").replace("#", "")[:600]
                sound = BytesIO()
                tts = gTTS(text=audio_text, lang='ro')
                tts.write_to_fp(sound)
                st.audio(sound, format='audio/mp3')

        except Exception as e:
            st.error(f"Eroare: {e}")
